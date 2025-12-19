"""
Document parsing service for various file formats including images with OCR.
"""
import io
import os
import uuid
import re
from typing import List, Optional
# import PyPDF2
import fitz
from pdfminer.high_level import extract_text as pdfminer_extract_text
from docx import Document
import pandas as pd
from PIL import Image, ImageOps, ImageFilter
import numpy as np
import pytesseract
from loguru import logger

from ..models import DocumentChunk
from ..config import settings
from .text_cleaner import text_cleaner

class DocumentParser:
    """Document parsing service for multiple file formats."""
    
    def __init__(self):
        """Initialize the document parser."""
        self.chunk_size = settings.chunk_size
        self.chunk_overlap = settings.chunk_overlap
        
        # Supported image formats for OCR
        self.image_formats = {'png', 'jpg', 'jpeg', 'tiff', 'tif', 'bmp', 'gif'}
        
        # OCR configuration
        self.ocr_config = r'--oem 3 --psm 6'  # Use LSTM OCR Engine Mode with uniform text block
        # Optional OCR language and whitelist from environment
        self.ocr_lang = os.getenv('TESSERACT_LANG', 'eng')
        self.ocr_char_whitelist = os.getenv('TESSERACT_CHAR_WHITELIST')

    def _preprocess_for_ocr(self, image: Image.Image) -> Image.Image:
        """Apply preprocessing recommended by Tesseract docs to improve OCR quality.

        Steps (inspired by tessdoc [ImproveQuality.md](https://github.com/tesseract-ocr/tessdoc/blob/main/ImproveQuality.md)):
        - Convert to grayscale
        - Scale up toward ~300 DPI equivalent (heuristic by upscaling small images)
        - Auto-contrast to stretch histogram
        - Median filter to reduce salt-and-pepper noise
        - Otsu binarization
        - Sharpen lightly
        - Deskew using Tesseract OSD angle when reliably detected
        """
        # Normalize transparency by compositing over white if present
        if image.mode in ('RGBA', 'LA') or (hasattr(image, 'info') and image.info.get('transparency') is not None):
            base = Image.new('RGBA', image.size, (255, 255, 255, 255))
            image = Image.alpha_composite(base, image.convert('RGBA')).convert('RGB')

        # Convert to grayscale
        gray = image.convert('L')

        # Heuristic upscaling for low-res images (target small side >= 1200 px)
        min_side = min(gray.size)
        if min_side < 1200:
            scale = 1200 / float(min_side)
            new_size = (int(gray.width * scale), int(gray.height * scale))
            gray = gray.resize(new_size, Image.LANCZOS)

        # Auto-contrast
        gray = ImageOps.autocontrast(gray)

        # Median denoise
        gray = gray.filter(ImageFilter.MedianFilter(size=3))

        # Otsu binarization using numpy
        np_img = np.asarray(gray)
        # Compute histogram and between-class variance to get Otsu threshold
        hist, bin_edges = np.histogram(np_img.flatten(), bins=256, range=(0, 255))
        total = np_img.size
        sum_total = np.dot(np.arange(256), hist)
        sumB = 0.0
        wB = 0.0
        max_var = 0.0
        threshold = 127
        for t in range(256):
            wB += hist[t]
            if wB == 0:
                continue
            wF = total - wB
            if wF == 0:
                break
            sumB += t * hist[t]
            mB = sumB / wB
            mF = (sum_total - sumB) / wF
            between_var = wB * wF * (mB - mF) ** 2
            if between_var > max_var:
                max_var = between_var
                threshold = t
        binary = (np_img > threshold).astype(np.uint8) * 255
        # Ensure white background and black text (invert if needed)
        num_white = int((binary == 255).sum())
        num_black = int((binary == 0).sum())
        if num_black > num_white:
            binary = 255 - binary
        bin_img = Image.fromarray(binary, mode='L')

        # Light sharpen to improve edge clarity after binarization
        bin_img = bin_img.filter(ImageFilter.UnsharpMask(radius=1, percent=150, threshold=3))

        # Deskew using Tesseract's OSD
        try:
            angle = self._estimate_skew_angle(bin_img)
            # Only rotate if angle is significant
            if angle and abs(angle) >= 0.5:
                # Expand True to avoid cropping
                bin_img = bin_img.rotate(-angle, resample=Image.BICUBIC, expand=True, fillcolor=255)
        except Exception as e:
            logger.warning(f"Deskew skipped due to OSD error: {e}")
        
        return bin_img

    def _estimate_skew_angle(self, image: Image.Image) -> float:
        """Estimate skew angle using pytesseract OSD.

        Returns angle in degrees, positive for counter-clockwise.
        """
        try:
            osd = pytesseract.image_to_osd(image)
            # Typical line in OSD output: 'Rotate: 270\n'
            match = re.search(r"Rotate:\s*(\d+)", osd)
            if match:
                angle = int(match.group(1))
                # Tesseract returns one of {0, 90, 180, 270}; keep as float
                return float(angle if angle != 270 else -90)
        except Exception as e:
            logger.debug(f"OSD angle detection failed: {e}")
        return 0.0
    
    async def parse_document(
        self, 
        file_content: bytes, 
        file_type: str, 
        filename: str
    ) -> Optional[List[DocumentChunk]]:
        """
        Parse a document and return text chunks.
        
        Args:
            file_content: Raw file content as bytes
            file_type: Type of the file (pdf, docx, txt, csv, png, jpg, jpeg, tiff, tif, bmp, gif)
            filename: Original filename for logging
            
        Returns:
            List of DocumentChunk objects, or None if parsing fails
        """
        try:
            logger.info(f"Starting to parse {file_type.upper()} file: {filename}")
            
            # Extract text based on file type
            if file_type.lower() == 'txt':
                text = await self._parse_txt(file_content)
            elif file_type.lower() == 'pdf':
                text = await self._parse_pdf(file_content)
            elif file_type.lower() == 'docx':
                text = await self._parse_docx(file_content)
            elif file_type.lower() == 'csv':
                text = await self._parse_csv(file_content)
            elif file_type.lower() in self.image_formats:
                text = await self._parse_image(file_content, filename)
            else:
                logger.error(f"Unsupported file type: {file_type}")
                return None
            
            if not text or not text.strip():
                logger.warning(f"No text extracted from file: {filename}")
                return []
            
            # Clean the extracted text before chunking
            logger.info(f"Cleaning extracted text from {filename}")
            cleaned_text = text_cleaner.clean_document_text(text, aggressive=False)
            
            if not cleaned_text or not cleaned_text.strip():
                logger.warning(f"No text remaining after cleaning: {filename}")
                return []
            
            # Log cleaning statistics
            stats = text_cleaner.get_cleaning_stats(text, cleaned_text)
            logger.info(f"Text cleaning stats for {filename}: {stats}")
            
            # Chunk the cleaned text
            chunks = await self._chunk_text(cleaned_text, filename)
            logger.info(f"Successfully parsed {filename} into {len(chunks)} chunks")
            return chunks
            
        except Exception as e:
            logger.error(f"Error parsing document {filename}: {e}")
            return None
    
    async def _parse_txt(self, file_content: bytes) -> str:
        """Parse TXT file."""
        try:
            # Try UTF-8 first, then fallback to other encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    logger.info(f"Successfully decoded TXT file with {encoding} encoding")
                    return text
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, use utf-8 with error handling
            text = file_content.decode('utf-8', errors='replace')
            logger.warning("Used UTF-8 with error replacement for TXT file")
            return text
            
        except Exception as e:
            logger.error(f"Error parsing TXT file: {e}")
            raise
    
    def _recover_pix(self, doc, item):
        xref = item[0]  # xref of PDF image
        smask = item[1]  # xref of its /SMask

        # special case: /SMask or /Mask exists
        if smask > 0:
            pix0 = fitz.Pixmap(doc.extract_image(xref)["image"])
            if pix0.alpha:  # catch irregular situation
                pix0 = fitz.Pixmap(pix0, 0)  # remove alpha channel
            mask = fitz.Pixmap(doc.extract_image(smask)["image"])

            try:
                pix = fitz.Pixmap(pix0, mask)
            except Exception:  # fallback to original base image in case of problems
                pix = fitz.Pixmap(doc.extract_image(xref)["image"])

            if pix0.n > 3:
                ext = "pam"
            else:
                ext = "png"

            return {  # create dictionary expected by caller
                "ext": ext,
                "colorspace": pix.colorspace.n,
                "image": pix.tobytes(ext),
            }

        # special case: /ColorSpace definition exists
        # to be sure, we convert these cases to RGB PNG images
        if "/ColorSpace" in doc.xref_object(xref, compressed=True):
            pix = fitz.Pixmap(doc, xref)
            pix = fitz.Pixmap(fitz.csRGB, pix)
            return {  # create dictionary expected by caller
                "ext": "png",
                "colorspace": 3,
                "image": pix.tobytes("png"),
            }

        return doc.extract_image(xref)

    async def _parse_pdf(self, file_content: bytes) -> str:
        """Parse PDF file using multiple methods."""
        text = ""

        try:
            pdf_reader = fitz.open(stream=file_content, filetype="pdf")

            for page_num, page in enumerate(pdf_reader):
                blocks = []

                # --- TEXT BLOCKS: use "dict" for richer structure ---
                for block in page.get_text("dict")["blocks"]:
                    if "lines" in block:
                        block_text = ""
                        for line in block["lines"]:
                            for span in line["spans"]:
                                block_text += span.get("text", "")
                        blocks.append({
                            "type": "text",
                            "bbox": block["bbox"],
                            "content": block_text.strip()
                        })

                # FIXME: XREF masih NONE.
                for block in page.get_text("rawdict")["blocks"]:
                    if block.get("type") == 1 or "image" in block:
                        extracted_content = await self._parse_image(block.get("image"), f"page{page_num}_img")
                        blocks.append({
                            "type": "image",
                            "bbox": block["bbox"],
                            "content": extracted_content
                        })

                blocks = sorted(blocks, key=lambda b: (b["bbox"][1], b["bbox"][0]))
                text += "\n".join(
                    b["content"] for b in blocks if b.get("content")
                )
            
            logger.info(f"Successfully extracted {len(text)} characters from PDF")
            return text
            
        except Exception as e:
            logger.error(f"Error parsing PDF file: {e}")
            # Last resort: try pdfminer if PyPDF2 completely failed
            try:
                file_stream = io.BytesIO(file_content)
                text = pdfminer_extract_text(file_stream)
                logger.info("Fallback to pdfminer succeeded")
                return text
            except Exception as fallback_error:
                logger.error(f"Pdfminer fallback also failed: {fallback_error}")
                raise
    
    async def _parse_docx(self, file_content: bytes) -> str:
        """Parse DOCX file."""
        try:
            file_stream = io.BytesIO(file_content)
            doc = Document(file_stream)
            
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            text = "\n".join(text_parts)
            logger.info(f"Successfully extracted {len(text)} characters from DOCX")
            return text
            
        except Exception as e:
            logger.error(f"Error parsing DOCX file: {e}")
            raise
    
    async def _parse_csv(self, file_content: bytes) -> str:
        """Parse CSV file and convert to readable text format."""
        try:
            # Try different encodings for CSV
            encodings = ['utf-8', 'latin-1', 'cp1252']
            df = None
            
            for encoding in encodings:
                try:
                    file_stream = io.StringIO(file_content.decode(encoding))
                    df = pd.read_csv(file_stream)
                    logger.info(f"Successfully read CSV with {encoding} encoding")
                    break
                except (UnicodeDecodeError, pd.errors.EmptyDataError):
                    continue
            
            if df is None:
                raise ValueError("Could not read CSV file with any encoding")
            
            # Convert DataFrame to readable text format
            text_parts = []
            
            # Add column headers
            text_parts.append("Column Headers: " + ", ".join(df.columns.tolist()))
            text_parts.append("")  # Empty line
            
            # Add data summary
            text_parts.append(f"Total Rows: {len(df)}")
            text_parts.append(f"Total Columns: {len(df.columns)}")
            text_parts.append("")
            
            # Add sample data (first few rows)
            sample_size = min(10, len(df))  # Show first 10 rows or less
            text_parts.append(f"Sample Data (first {sample_size} rows):")
            
            for idx, row in df.head(sample_size).iterrows():
                row_text = f"Row {idx + 1}: "
                row_values = []
                for col, value in row.items():
                    if pd.notna(value):  # Skip NaN values
                        row_values.append(f"{col}: {value}")
                row_text += ", ".join(row_values)
                text_parts.append(row_text)
            
            # Add column statistics for numeric columns
            numeric_columns = df.select_dtypes(include=['number']).columns
            if len(numeric_columns) > 0:
                text_parts.append("")
                text_parts.append("Numeric Column Statistics:")
                for col in numeric_columns:
                    stats = df[col].describe()
                    text_parts.append(
                        f"{col}: Mean={stats['mean']:.2f}, "
                        f"Min={stats['min']:.2f}, Max={stats['max']:.2f}"
                    )
            
            text = "\n".join(text_parts)
            logger.info(f"Successfully converted CSV to text ({len(text)} characters)")
            return text
            
        except Exception as e:
            logger.error(f"Error parsing CSV file: {e}")
            raise
    
    async def _parse_image(self, file_content: bytes, filename: str) -> str:
        """Parse image file using OCR to extract text."""
        try:
            logger.info(f"Starting OCR processing for image: {filename}")
            
            # Open image using PIL
            image_stream = io.BytesIO(file_content)
            image = Image.open(image_stream)
            
            # Convert to RGB if necessary (some formats might be in different modes)
            if image.mode != 'RGB':
                logger.info(f"Converting image from {image.mode} to RGB")
                image = image.convert('RGB')
            
            # Log image properties
            logger.info(f"Image properties - Size: {image.size}, Mode: {image.mode}, Format: {image.format}")
            
            # Preprocess image per tessdoc recommendations
            preprocessed = self._preprocess_for_ocr(image)
            logger.info(f"Preprocessed image size: {preprocessed.size}")

            # Build OCR config with optional language and whitelist
            base_config = self.ocr_config
            if self.ocr_char_whitelist:
                base_config = base_config + f" -c tessedit_char_whitelist={self.ocr_char_whitelist}"
            lang = self.ocr_lang or 'eng'

            # Perform OCR using pytesseract
            try:
                # First try with default configuration
                text = pytesseract.image_to_string(preprocessed, lang=lang, config=base_config)
                logger.info(f"OCR extracted {len(text)} characters with default config")
                
                # If we got very little text, try with different PSM modes
                if len(text.strip()) < 50:
                    logger.info("Trying alternative OCR configurations for better text extraction")
                    
                    # Try different Page Segmentation Modes
                    alternative_configs = [
                        r'--oem 3 --psm 3',  # Fully automatic page segmentation
                        r'--oem 3 --psm 4',  # Assume a single column of text
                        r'--oem 3 --psm 7',  # Treat as a single text line
                        r'--oem 3 --psm 8',  # Treat as a single word
                        r'--oem 3 --psm 11', # Sparse text
                        r'--oem 3 --psm 13'  # Raw line
                    ]
                    
                    best_text = text
                    for config in alternative_configs:
                        try:
                            cfg = config
                            if self.ocr_char_whitelist:
                                cfg = cfg + f" -c tessedit_char_whitelist={self.ocr_char_whitelist}"
                            alt_text = pytesseract.image_to_string(preprocessed, lang=lang, config=cfg)
                            if len(alt_text.strip()) > len(best_text.strip()):
                                best_text = alt_text
                                logger.info(f"Better OCR result with config: {config}")
                        except Exception as config_error:
                            logger.warning(f"OCR config {config} failed: {config_error}")
                    
                    text = best_text
                
            except Exception as ocr_error:
                logger.error(f"OCR processing failed: {ocr_error}")
                # Try with minimal configuration as fallback
                try:
                    text = pytesseract.image_to_string(preprocessed, lang=lang)
                    logger.info("Fallback OCR with minimal config succeeded")
                except Exception as fallback_error:
                    logger.error(f"Fallback OCR also failed: {fallback_error}")
                    raise
            
            # Clean up the OCR text
            if text:
                # Remove common OCR artifacts
                text = self._clean_ocr_text(text)
                logger.info(f"Cleaned OCR text length: {len(text)}")
            
            logger.info(f"Successfully extracted text from image {filename}")
            return text
            
        except Exception as e:
            logger.error(f"Error parsing image file {filename}: {e}")
            raise
    
    def _clean_ocr_text(self, text: str) -> str:
        """Clean OCR-specific artifacts from extracted text."""
        if not text:
            return text
        
        # Common OCR cleaning
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove standalone single characters that are likely OCR errors
        # but keep meaningful single characters like 'a', 'I'
        words = text.split()
        cleaned_words = []
        
        for word in words:
            # Skip single characters that are likely OCR noise
            if len(word) == 1 and word not in 'aAiI' and not word.isdigit():
                continue
            # Skip words that are mostly special characters
            if len(word) > 1 and sum(c in '!@#$%^&*()_+-=[]{}|;:,.<>?' for c in word) > len(word) * 0.7:
                continue
            cleaned_words.append(word)
        
        text = ' '.join(cleaned_words)
        
        # Remove common OCR misreads
        ocr_artifacts = [
            r'\b[Il1]\b(?=\s[a-z])',  # Standalone I, l, 1 followed by lowercase (likely OCR error)
            r'\b[oO0]\b(?=\s[a-z])',  # Standalone o, O, 0 followed by lowercase
        ]
        
        for pattern in ocr_artifacts:
            text = re.sub(pattern, '', text)
        
        # Final whitespace cleanup
        text = re.sub(r'\s+', ' ', text).strip()
        
        return text
    
    async def _chunk_text(self, text: str, filename: str) -> List[DocumentChunk]:
        """
        Split text into chunks with overlap.
        
        Args:
            text: Input text to chunk
            
        Returns:
            List of DocumentChunk objects
        """
        if not text or not text.strip():
            return []
        
        chunks = []
        text = text.strip()
        
        # Simple character-based chunking with overlap
        start = 0
        chunk_index = 0
        
        while start < len(text):
            # Calculate end position
            end = start + self.chunk_size
            
            # If this is not the last chunk, try to break at word boundary
            if end < len(text):
                # Look for the last space within the chunk to avoid breaking words
                last_space = text.rfind(' ', start, end)
                if last_space != -1 and last_space > start:
                    end = last_space
            else:
                end = len(text)
            
            # Extract chunk text
            chunk_text = text[start:end].strip()
            chunk_text = f"{filename} - Chunk {chunk_index}: {chunk_text}" # better capturing the context of the chunk
            
            if chunk_text:  # Only add non-empty chunks
                chunk = DocumentChunk(
                    text=chunk_text,
                    index=chunk_index,
                    start_char=start,
                    end_char=end
                )
                chunks.append(chunk)
                chunk_index += 1
            
            # Move start position with overlap
            if end >= len(text):
                break
            
            start = end - self.chunk_overlap
            # Ensure we don't go backwards
            if start <= chunks[-1].start_char if chunks else 0:
                start = end
        
        logger.info(f"Created {len(chunks)} chunks from text of length {len(text)}")
        return chunks


# Global parser instance
document_parser = DocumentParser()
